from flask import Flask, request, jsonify, make_response
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import io
import os
from datetime import datetime
from urllib.parse import quote

app = Flask(__name__, static_folder='static', static_url_path='')
CORS(app)

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    try:
        data = request.json
        
        client_name = data.get('clientName', '')
        execute_date = data.get('executeDate', '')
        weight = data.get('weight', '')
        photos = data.get('photos', [])
        attachments = data.get('attachments', [])
        
        # Convert date format
        if execute_date:
            date_obj = datetime.strptime(execute_date, '%Y-%m-%d')
            roc_year = date_obj.year - 1911
            formatted_date = f'{roc_year}年{date_obj.month}月{date_obj.day}日'
        else:
            formatted_date = ''
        
        # Create Word document
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        
        # Title
        title1 = doc.add_paragraph()
        title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = title1.add_run('大豐環保科技股份有限公司')
        run1.font.size = Pt(20)
        run1.font.bold = True
        
        title2 = doc.add_paragraph()
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = title2.add_run(f'{client_name}-報廢回饋資料')
        run2.font.size = Pt(18)
        run2.font.bold = True
        
        doc.add_paragraph()
        
        # Info table
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        cell = info_table.rows[0].cells[0]
        cell.text = '客戶名稱'
        cell = info_table.rows[0].cells[1]
        cell.text = client_name
        
        cell = info_table.rows[1].cells[0]
        cell.text = '執行日期'
        cell = info_table.rows[1].cells[1]
        cell.text = formatted_date
        
        cell = info_table.rows[2].cells[0]
        cell.text = '重量'
        cell = info_table.rows[2].cells[1]
        cell.text = f'{weight} kg'
        
        doc.add_paragraph()
        
        # Photos section
        if photos:
            section_title = doc.add_paragraph()
            section_title_run = section_title.add_run('待報廢照片')
            section_title_run.font.size = Pt(16)
            section_title_run.font.bold = True
            doc.add_paragraph()
            
            photo_width = Inches(3.67)
            photo_height = Inches(2.67)
            
            for photo in photos:
                if photo.get('description'):
                    desc_para = doc.add_paragraph(photo['description'])
                    desc_para.runs[0].font.size = Pt(10)
                
                try:
                    img_data = photo.get('image', '')
                    if img_data.startswith('data:image'):
                        base64_data = img_data.split(',')[1]
                        img_bytes = base64.b64decode(base64_data)
                        image_stream = io.BytesIO(img_bytes)
                        try:
                            doc.add_picture(image_stream, width=photo_width, height=photo_height)
                        except:
                            doc.add_picture(image_stream, width=photo_width)
                except Exception as e:
                    print(f'Error adding image: {e}')
                
                doc.add_paragraph()
        
        # Attachments section
        if attachments:
            section_title = doc.add_paragraph()
            section_title_run = section_title.add_run('附件資料')
            section_title_run.font.size = Pt(16)
            section_title_run.font.bold = True
            doc.add_paragraph()
            
            for attachment in attachments:
                if attachment.get('description'):
                    desc_para = doc.add_paragraph(attachment['description'])
                    desc_para.runs[0].font.size = Pt(10)
                
                try:
                    img_data = attachment.get('image', '')
                    if img_data.startswith('data:image'):
                        base64_data = img_data.split(',')[1]
                        img_bytes = base64.b64decode(base64_data)
                        image_stream = io.BytesIO(img_bytes)
                        doc.add_picture(image_stream, width=Inches(6.0))
                except Exception as e:
                    print(f'Error adding attachment: {e}')
                
                doc.add_paragraph()
        
        # Save to memory
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Generate filename
        filename = f'廢棄物清運報告_{client_name}.docx'
        filename_encoded = quote(filename)
        
        response = make_response(doc_stream.getvalue())
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{filename_encoded}"
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response
        
    except Exception as e:
        print(f'Error: {e}')
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    print(f'啟動伺服器於 http://localhost:{port}')
    app.run(host='0.0.0.0', port=port)
